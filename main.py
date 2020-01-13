import sqlite3
from tkinter import *
from database import *


class lecturer():
    def __init__(self,first_name,last_name,email,password,phone):
        self.first_name=first_name
        self.last_name=last_name
        self.email=email
        self.password=password
        self.phone=phone
        
    conn = sqlite3.connect('login.db', timeout=1)
    c = conn.cursor()
    valid  = False
    taken = False
    loggedIn = False
    userName = False
    
    def create_table(self):
        self.c = self.conn.cursor()
        self.c.execute('''CREATE TABLE IF NOT EXISTS lecturer (
                                            id INTEGER PRIMARY KEY AUTOINCREMENT,                                        
                                            first_name text NOT NULL,
                                            last_name text NOT NULL,
                                            email TEXT NOT NULL,
                                            password TEXT,
                                            phone TEXT NOT NULL 
                                        );''')
        
    
    def data_entry(self,first_name, last_name,email,password,phone):
        self.c = self.conn.cursor()
        self.c.execute("INSERT INTO lecturer (first_name, last_name,email,password,phone) VALUES (?, ?, ? , ?, ?)",
                  (first_name, last_name,email,password,phone))
        access='1'
        self.c.execute("INSERT INTO users (email, password,access) VALUES (?, ?, ?)",
                  (email, password,access))        
        self.conn.commit()
        #self.c.close()
        #self.conn.close()
    
    def register(self):
        self.c = self.conn.cursor()
        valid  = False
        taken = False

        while not valid:
            self.c.execute("SELECT * FROM lecturer")
            taken = False            
            for row in self.c.fetchall():
                if str(self.email) == str(row[4]):
                    print("that email is already being used")
                    taken = True                  
            if not taken:
                valid = True            
        self.data_entry(self.first_name,self.last_name,self.email,self.password,self.phone)
        
    def removeLec(self):
        self.c = self.conn.cursor()
        self.c.execute("""select * from lecturer""")
        allList=self.c.fetchone()
        for row in allList:
            print(allList)
        
        choise=input('Enter the number lecturer u want remove: ')
        self.c.execute('''DELETE FROM lecturer
                  WHERE id = ?;''',choise)
        self.conn.commit()
        print('remove sucessed')
        #self.c.close()
        #self.conn.close()
        
    def changeLec(self):
        self.c = self.conn.cursor()
        self.c.execute("""select * from lecturer""")
        allList=self.c.fetchone()            
        for row in allList:
            print(allList)
        whatUwantChange=input('what u want to change\n1.first name\n2.last name\n3.email\n4.password\n5.phone')   
        idChose=input('Enter the id of workers u want to change ') #איזה מספר עובד ברצונך לשנות
        if whatUwantChange == 1:
            update=input('Enter the new update')
            self.c.execute('''UPDATE lecturer
                              SET first_name = ?
                              WHERE id = ?;''',update,idChose)
            print('change done!')
            self.conn.commit()
        if whatUwantChange == 2:
            update=input('Enter the new update')
            self.c.execute('''UPDATE lecturer
                              SET last_name = ?
                              WHERE id = ?;''',update,idChose)
            self.conn.commit()  
        if whatUwantChange == 3:
            update=input('Enter the new update')
            self.c.execute('''UPDATE lecturer
                              SET email = ?
                              WHERE id = ?;''',update,idChose)
            self.conn.commit()  
        if whatUwantChange == 4:
            update=input('Enter the new update')
            self.c.execute('''UPDATE lecturer
                              SET password = ?
                              WHERE id = ?;''',update,idChose)
            self.conn.commit()
        if whatUwantChange == 5:
            update=input('Enter the new update')
            self.c.execute('''UPDATE lecturer
                              SET phone = ?
                              WHERE id = ?;''',update,idChose)
            self.conn.commit()  
class student():
    def __init__(self,first_name,last_name,email,password,phone):
        self.first_name=first_name
        self.last_name=last_name
        self.email=email
        self.password=password
        self.phone=phone
        
    conn = sqlite3.connect('login.db', timeout=1)
    c = conn.cursor()
    valid  = False
    taken = False
    loggedIn = False
    userName = False
    
    def create_table(self):
        self.c = self.conn.cursor()
        self.c.execute('''CREATE TABLE IF NOT EXISTS student (
                                            id INTEGER PRIMARY KEY AUTOINCREMENT,                                        
                                            first_name text NOT NULL,
                                            last_name text NOT NULL,
                                            email TEXT NOT NULL,
                                            password TEXT,
                                            phone TEXT NOT NULL 
                                        );''')
        
    
    def data_entry(self,first_name, last_name,email,password,phone):
        self.c = self.conn.cursor()
        self.c.execute("INSERT INTO student (first_name, last_name,email,password,phone) VALUES (?, ?, ? , ?, ?)",
                  (first_name, last_name,email,password,phone))
        access='1'
        self.c.execute("INSERT INTO users (email, password,access) VALUES (?, ?, ?)",
                  (email, password,access))        
        self.conn.commit()
        #self.c.close()
        #self.conn.close()
    
    def register(self):
        self.c = self.conn.cursor()
        valid  = False
        taken = False

        while not valid:
            self.c.execute("SELECT * FROM student")
            taken = False            
            for row in self.c.fetchall():
                if str(self.email) == str(row[4]):
                    print("that email is already being used")
                    taken = True                  
            if not taken:
                valid = True            
        self.data_entry(self.first_name,self.last_name,self.email,self.password,self.phone)
        
    def removeStu(self):
        self.c = self.conn.cursor()
        self.c.execute("""select * from student""")
        allList=self.c.fetchone()
        for row in allList:
            print(allList)
        
        choise=input('Enter the number lecturer u want remove: ')
        self.c.execute('''DELETE FROM student
                  WHERE id = ?;''',choise)
        self.conn.commit()
        print('remove sucessed')
        #self.c.close()
        #self.conn.close()
class Course():
    def __init__(self,name,lastname,coursename, year,semester, season):        
        self.name=name
        self.lastname=lastname
        self.coursename=coursename
        self.year=year
        self.semester=semester
        self.season=season
        
    def CreatFile(self):
        data={}        
        data['Course']=[]        
        data['Course'].append({"name:":self.name , "last_name" :self.lastname , "coursename" : self.coursename ,
                         'year':self.year,'semester':self.semester,'season':self.season })
        #self.number=number+1
        import json       
        with open('Course.txt', 'w') as outfile:
            json.dump(data , outfile)
            outfile.write("\n")
    def addToFile(self):
        data={}        
        data['Course']=[]
        data['Course'].append({"name:":self.name , "last_name" :self.lastname , "coursename" : self.coursename ,
                         'year':self.year,'semester':self.semester,'season':self.season })
        #self.number=number+1
        import json       
        with open('Course.txt', 'a') as outfile:
            json.dump(data , outfile)
            outfile.write("\n")  
            
            
class Questions(Course):
    def __init__(self,Subject,SubSubject,level,Type,year,semester,mued,Format,Solution):
       self.Subject=Subject
       self.SubSubject=SubSubject
       self.level=level
       self.Type=Type
       self.year=year
       self.semester=semester
       self.mued=mued
       self.Format=Format
       self.Solution= Solution
       self.number=1

    def CreatFile(self):
        data={}        
        data['Quest']=[]
        data['Quest'].append({"number:":self.number , "Subject" :self.Subject , "SubSubject" : self.SubSubject ,
                         'level':self.level,'Type':self.Type,'year':self.year,'semester':self.semester,                       
                         'mued':self.mued,'Format':self.Format,'Solution':self.Solution })
        #self.number=number+1
        import json       
        with open('Questions.txt', 'w') as outfile:
            json.dump(data , outfile)
            outfile.write("\n")
    def addToFile(self):
        data={}        
        data['Quest']=[]
        data['Quest'].append({"number:":self.number , "Subject" :self.Subject , "SubSubject" : self.SubSubject ,
                         'level':self.level,'Type':self.Type,'year':self.year,'semester':self.semester,                       
                         'mued':self.mued,'Format':self.Format,'Solution':self.Solution})
        #self.number=number+1
        import json       
        with open('Questions.txt', 'a') as outfile:
            json.dump(data , outfile)
            outfile.write("\n") 
            
        


db = Database() #קריאה למחלקה בקובץ database
main_window = Tk()
main_window.title('Team Program')
main_window.geometry('300x200')
loginFrame = Frame(main_window)
loginFrame.pack()

Label(loginFrame, justify='left', text='Hello,\nplease enter your email and password below.', fg='black').pack()
statusMessage = Label(loginFrame)
statusMessage.pack()


email = Entry(loginFrame)
password = Entry(loginFrame, show='*')

Label(loginFrame, text='email:').pack()
email.pack()
Label(loginFrame, text='Password:').pack()
password.pack()

email.focus();

coordOption = IntVar()
coordOption.set(0)

def coordinatorH():
    print(coordOption.get())

LecturerOption = IntVar()
LecturerOption.set(0)

def coordinatorH():
    print(coordOption.get())

LecturerOption = IntVar()
LecturerOption.set(0)
def LecturerH():
    print(LecturerOption.get())

def tryLogin(event=0):
    print(email.get() + " " + password.get())
    query = db.execute('SELECT * FROM users WHERE email="%s" and password="%s"'%(email.get(), password.get()))
    user = query.fetchone()
    if user:
        statusMessage.configure(text='You are successfully logged-in!', fg='green')
        loginFrame.destroy()

        main_window.title('Logged in as %s(%s)'%(user[1], user[3]))
        loggedInFrame = Frame(main_window)

        str = "Hello %s," % user[1]
        Label(loggedInFrame, text=str, font=('Arial', 25), anchor='nw', width=50).pack()
        if user[3]=='student':
            pass
        elif user[3]=='lecturer':
            main_window.geometry('500x500')

            Label(loggedInFrame, text="Choose what action you want to take:",anchor='w',width=100,fg='green').pack()
            for i in range(len(LecturerMenu)):
                Radiobutton(loggedInFrame, text=LecturerMenu[i], anchor='w', width=100, variable=LecturerOption,value=i).pack()
            Button(loggedInFrame, text='OK', width=17, command=LecturerH).pack(pady=40)
        elif user[3]=='coordinator':
            main_window.geometry('500x500')

            Label(loggedInFrame, text="Choose what action you want to take:",anchor='w',width=100,fg='green').pack()
            for i in range(len(CoordinatorMenu)):
                Radiobutton(loggedInFrame, text=CoordinatorMenu[i], anchor='w', width=100, variable=coordOption, value=i).pack()

            Button(loggedInFrame, text='OK', width=17, command=coordinatorH).pack(pady=40)

        loggedInFrame.pack()
    else:
        statusMessage.configure(text='email or password are invalid!',fg='red')
    query.close()
loginButton = Button(loginFrame, text='Log me in', width=17, command=tryLogin)
loginButton.pack(pady=9)
main_window.bind('<Return>', tryLogin)
main_window.mainloop()


            
def addLecturer():     #להוסיף מרצה       
    first_name=input('Enter first name : ')
    last_name = input('Enter last name : ')
    email= input('Enter your email : ')
    password = input('Enter password for your account : ')
    phone = input('Enter phone number : ')
    global a
    a=lecturer(first_name,last_name,email,password,phone)
    a.create_table()
    a.register()

def removeLecturer():
    a.removeLec()
    
def addStudent():     #להוסיף סטודנט       
    first_name=input('Enter first name : ')
    last_name = input('Enter last name : ')
    email= input('Enter your email : ')
    password = input('Enter password for your account : ')
    phone = input('Enter phone number : ')
    global b
    b=student(first_name,last_name,email,password,phone)
    b.create_table()
    b.register()

def removeStudent():
    b.removeStu()
    

def menuOfCoordinator():
    stop=True
    while stop:
        print('Coordinator menu')
        print('What u want to do ?')
        print('Press 1 to Add Lecturer')
        print('Press 2 to Add Student')
        print('Press 3 to remove Lecturer')
        print('Press 4 to remove Student')
        print('Press 5 to Cut from PDF')
        print('Press 6 to exit from coordinator menu')
        x=int(input('Enter now: '))
        if x==1:
            addLecturer()
            break
        if x==2:
            addStudent()
            break
        if x==3:
            removeLecturer()
        if x==4:
            removeStudent()
        if x==5:
            print('the program open web site to Convert PDF to JPG,save the jpg in project file')
            import webbrowser
            webbrowser.open('https://smallpdf.com/pdf-to-jpg')
            print('after you Convert PDF to JPG cut in next site the Question you want to past!')
            print('Save as 0001.jpg')
            webbrowser.open('https://www.iloveimg.com/crop-image/crop-jpg')

				
            from docx import Document
            from docx.shared import Inches

            document = Document()

            p = document.add_paragraph()
            r = p.add_run()
            r.add_picture('0001.jpg')            
            document.save('image.docx')
            break
        if x==6:
            stop=False
            break
        else:
            print('Worng Choise,Try Again\n')
            continue
        
def menuOfLecturer():
    stop = True
    while stop:
        print('Lecturer menu')
        print('What u want to do ?')
        print('Press 1 to CreatFile for Course')
        print('Press 2 to CreatFile for Quest')
        print('Press 3 to add a Course')
        print('Press 4 to add a Quest')
        print('Press 5 to exit from Lecturer menu')

        x=int(input('Enter now: '))
        if x==1:
            choiseCourse=int(input('how many Course u want?'))
            NumberOfCourse=[]
            for i in range(choiseCourse):
                print(f'Course number {i+1}..')
                name= input('Enter a name: ')
                lastname=input('Enter a lastname: ')
                coursename=input('Enter a coursename: ')
                year=input('Enter a year: ')
                semester=input('Enter a semester: ')
                season=input('Enter a season: ')
                NumberOfCourse.append(Course(name,lastname,coursename,year,semester,season))
            NumberOfCourse[0].CreatFile()
            break
        if x==2:            
            choiseQuest=int(input('how many Quest u want?'))
            NumberOfQuest=[]
            for i in range(choiseQuest):
                Subject=input('Enter a Subject: ')
                SubSubject= input('Enter a Sub Subject: ')
                level=input('Enter a level: ')
                Type=input('Enter a type: ')
                year= input('Enter a year: ')
                semester= input('Enter a semester: ')
                mued=input('Enter a mued: ')
                Format=input('Enter a Format: ')
                Solution=input('Enter a Solution: ')
                NumberOfQuest.append(Questions(Subject,SubSubject,level,Type,year,semester,mued,Format,Solution))
            NumberOfQuest[0].CreatFile()
            break
        if x==3:
            print('add A Course now..\n')
            NumberOfQuest=[]
            Subject=input('Enter a Subject: ')
            SubSubject= input('Enter a Sub Subject: ')
            level=input('Enter a level: ')
            Type=input('Enter a type: ')
            year= input('Enter a year: ')
            semester= input('Enter a semester: ')
            mued=input('Enter a mued: ')
            Format=input('Enter a Format: ')
            Solution=input('Enter a Solution: ')
            NumberOfCourse.append(Course(Subject,SubSubject,level,Type,year,semester,mued,Format,Solution))
            NumberOfCourse[i].addToFile()
            break
        if x==4:
            print('add A questions now..\n')
            NumberOfQuest=[]
            Subject=input('Enter a Subject: ')
            SubSubject= input('Enter a Sub Subject: ')
            level=input('Enter a level: ')
            Type=input('Enter a type: ')
            year= input('Enter a year: ')
            semester= input('Enter a semester: ')
            mued=input('Enter a mued: ')
            Format=input('Enter a Format: ')
            Solution=input('Enter a Solution: ')
            NumberOfQuest.append(Questions(Subject,SubSubject,level,Type,year,semester,mued,Format,Solution))
            NumberOfQuest[0].addToFile()
            break
        if x==5:
            stop=False
            break
        else:
            print('Worng Choise,Try Again\n')
            continue
            
    
menuOfCoordinator()
menuOfLecturer()
